from pathlib import Path
from docx import Document
import subprocess, os, shutil, zipfile

BASE = Path('/mnt/data/update_psychology/theory-site-static')
FILES = BASE/'files'

old = "לפני שהפרק נכנס אל המודל עצמו, צריך להבהיר את גבולו: שלוש הקומות שיופיעו כאן אינן שלוש האמיתות היחידות של הקיום. הן שלוש שפות מקבילות מתוך אין־ספור שפות אפשריות שבהן אפשר לתאר את אותו חוק רקורסיבי. אפשר היה להעביר את אותה תנועה דרך מבנה של מדינה: ריבונות, מוסדות, חוקים, אזרחים, גבולות, משבר אמון ותיקון חוקתי. אפשר היה להעביר אותה דרך כלכלה: מטבע, חוב, אשראי, אינפלציה, שוק, עבודה, מחסור, עודף ואיזון. אפשר היה גם לפתח אותה דרך משפט, משפחה, עיר, מערכת אקולוגית, מוזיקה או גוף פוליטי. כל אחת מן השפות הללו יכולה לחשוף היבט אחר של אותו מבנה."

new = "לפני שהפרק נכנס אל המודל עצמו, צריך להבהיר את גבולו: שלוש הקומות שיופיעו כאן אינן שלוש האמיתות היחידות של הקיום. הן שלוש שפות מקבילות מתוך אין־ספור שפות אפשריות שבהן אפשר לתאר את אותו חוק רקורסיבי. אפשר היה להעביר את אותה תנועה דרך מבנה של מדינה: ריבונות, מוסדות, חוקים, אזרחים, גבולות, משבר אמון ותיקון חוקתי. אפשר היה להעביר אותה דרך כלכלה: מטבע, חוב, אשראי, אינפלציה, שוק, עבודה, מחסור, עודף ואיזון. אפשר היה גם להעביר אותה דרך פסיכולוגיה (כהבהרה: אף שהטקסט שלהלן נוגע מעט במישור הפסיכולוגי, הוא אינו מרחיב את הנושא ואינו מפתח דוגמאות כגון שלבי התובנה בחייו של אדם: רגעים שבהם הוא מגיע להבנה או לאמת שנכונה לו באותו רגע, מתפרק ממנה, מתנסח מחדש, ונולד אל אפשרות מדויקת יותר - מעין גלגולי נשמות של תודעתו בתוך חייו). אפשר היה גם לפתח אותה דרך משפט, משפחה, עיר, מערכת אקולוגית, מוזיקה או גוף פוליטי. כל אחת מן השפות הללו יכולה לחשוף היבט אחר של אותו מבנה."

text_files = [
    FILES/'between-potential-and-ideal-he.md',
    FILES/'between-potential-and-ideal-he-editorial.html',
    FILES/'editorial-tightened/between-potential-and-ideal-tightened-he.md',
    FILES/'editorial-tightened/between-potential-and-ideal-tightened-he.html',
]
for p in text_files:
    s = p.read_text(encoding='utf-8')
    if old not in s:
        raise RuntimeError(f'old paragraph not found in {p}')
    p.write_text(s.replace(old, new, 1), encoding='utf-8')
    print('updated text', p)

for p in [FILES/'between-potential-and-ideal-he.docx', FILES/'editorial-tightened/between-potential-and-ideal-tightened-he.docx']:
    doc = Document(p)
    found = 0
    for para in doc.paragraphs:
        if para.text == old:
            # Keep original paragraph formatting and run-level style; replace the first run text.
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ''
            else:
                para.add_run(new)
            found += 1
    if found != 1:
        raise RuntimeError(f'expected 1 matching paragraph in {p}, found {found}')
    doc.save(p)
    print('updated docx', p)

# Convert updated Hebrew DOCX files to the linked PDFs.
convert_jobs = [
    (FILES/'between-potential-and-ideal-he.docx', FILES/'between-potential-and-ideal-he-editorial.pdf'),
    (FILES/'editorial-tightened/between-potential-and-ideal-tightened-he.docx', FILES/'editorial-tightened/between-potential-and-ideal-tightened-he.pdf'),
]
for docx_path, pdf_path in convert_jobs:
    out_dir = pdf_path.parent
    expected = out_dir/(docx_path.stem + '.pdf')
    if expected.exists():
        expected.unlink()
    env = os.environ.copy()
    env['HOME'] = '/tmp/lohome_psych_note'
    Path(env['HOME']).mkdir(exist_ok=True)
    res = subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir',str(out_dir),str(docx_path)], env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if res.returncode != 0:
        print(res.stdout)
        print(res.stderr)
        raise RuntimeError(f'LibreOffice conversion failed for {docx_path}')
    if not expected.exists():
        raise RuntimeError(f'Expected PDF not created: {expected}')
    if expected != pdf_path:
        if pdf_path.exists():
            pdf_path.unlink()
        expected.rename(pdf_path)
    print('updated pdf', pdf_path)

# zip all files
out_zip = Path('/mnt/data/theory_site_recursive_chapter_psychology_note_fixed.zip')
if out_zip.exists():
    out_zip.unlink()
with zipfile.ZipFile(out_zip, 'w', zipfile.ZIP_DEFLATED) as z:
    for f in BASE.parent.rglob('*'):
        if f.is_file():
            z.write(f, f.relative_to(BASE.parent))
print(out_zip)

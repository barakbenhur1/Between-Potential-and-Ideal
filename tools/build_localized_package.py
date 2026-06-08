#!/usr/bin/env python3
from __future__ import annotations

from argparse import ArgumentParser
from pathlib import Path
import html

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown import markdown
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
PUBLIC_FILES_ROOT = (ROOT / "site" / "files").resolve()
ALLOWED_LANGUAGES = {"tlh", "qya"}
ALLOWED_STATUSES = {"draft", "linguistic-review", "content-review", "approved"}
FONT_CANDIDATES = [
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
]


def parse_source(path: Path) -> tuple[dict[str, str], str]:
    text = path.read_text(encoding="utf-8")
    if not text.startswith("---\n"):
        raise SystemExit(f"Missing front matter in {path}")
    try:
        _, raw_meta, body = text.split("---\n", 2)
    except ValueError as exc:
        raise SystemExit(f"Malformed front matter in {path}") from exc

    meta: dict[str, str] = {}
    for line in raw_meta.splitlines():
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        key, sep, value = line.partition(":")
        if not sep:
            raise SystemExit(f"Invalid front matter line: {line}")
        meta[key.strip()] = value.strip()

    required = (
        "title",
        "language",
        "language_label",
        "direction",
        "status",
        "source_document",
        "semantic_cross_check",
    )
    for key in required:
        if not meta.get(key):
            raise SystemExit(f"Missing front matter key: {key}")

    if meta["language"] not in ALLOWED_LANGUAGES:
        raise SystemExit(f"Unsupported target language: {meta['language']}")
    if meta["direction"] not in {"ltr", "rtl"}:
        raise SystemExit("direction must be ltr or rtl")
    if meta["status"] not in ALLOWED_STATUSES:
        raise SystemExit(f"Unsupported translation status: {meta['status']}")
    if not body.strip():
        raise SystemExit(f"Translated body is empty: {path}")
    return meta, body.strip() + "\n"


def ensure_output_policy(meta: dict[str, str], output_stem: Path, allow_draft: bool) -> None:
    public_output = output_stem == PUBLIC_FILES_ROOT or PUBLIC_FILES_ROOT in output_stem.parents
    if public_output and meta["status"] != "approved":
        raise SystemExit(
            "Refusing to generate a public package from a non-approved source. "
            "Move the output outside site/files for review, or approve the source first."
        )
    if meta["status"] != "approved" and not allow_draft:
        raise SystemExit(
            f"Translation status is {meta['status']!r}; use --allow-draft only for private review output."
        )
    if public_output:
        expected_dir = (PUBLIC_FILES_ROOT / meta["language"]).resolve()
        if output_stem.parent != expected_dir:
            raise SystemExit(
                f"Public {meta['language']} packages must be written directly under "
                f"{expected_dir.relative_to(ROOT)}"
            )
        if not output_stem.name.endswith(f"-{meta['language']}"):
            raise SystemExit(
                f"Public package stem must end with -{meta['language']}: {output_stem.name}"
            )


def html_document(meta: dict[str, str], body: str) -> str:
    body_html = markdown(body, extensions=["extra", "sane_lists"])
    title = html.escape(meta["title"])
    lang = html.escape(meta["language"])
    direction = html.escape(meta["direction"])
    label = html.escape(meta["language_label"])
    return f"""<!doctype html>
<html lang="{lang}" dir="{direction}">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<meta name="robots" content="noindex" data-remove-when-language-published>
<title>{title} - Between Potential and Ideal</title>
<style>
body{{margin:0;background:#07090d;color:#eef2f6;font:18px/1.7 system-ui,sans-serif}}
main{{max-width:860px;margin:auto;padding:56px 24px 80px}}
.edition-note{{color:#9ce9dd;font-size:.9rem;margin-bottom:28px}}
h1,h2,h3{{color:#f3cf76;line-height:1.25}}
a{{color:#9ce9dd}} blockquote{{border-inline-start:4px solid #9f8cff;margin-inline:0;padding-inline:18px;color:#cbd3dc}}
</style>
</head>
<body><main><p class="edition-note">{label} edition</p>{body_html}</main></body>
</html>
"""


def plain_text(body: str) -> str:
    rendered = markdown(body, extensions=["extra", "sane_lists"])
    soup = BeautifulSoup(rendered, "html.parser")
    return soup.get_text("\n", strip=True) + "\n"


def build_docx(meta: dict[str, str], body: str, output: Path) -> None:
    rendered = markdown(body, extensions=["extra", "sane_lists"])
    soup = BeautifulSoup(rendered, "html.parser")
    doc = Document()
    alignment = WD_ALIGN_PARAGRAPH.RIGHT if meta["direction"] == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
    title = doc.add_heading(meta["title"], level=0)
    title.alignment = alignment
    for node in soup.find_all(recursive=False):
        name = node.name or ""
        text = node.get_text(" ", strip=True)
        if not text:
            continue
        if name in {"h1", "h2", "h3"}:
            level = {"h1": 1, "h2": 2, "h3": 3}[name]
            paragraph = doc.add_heading(text, level=level)
        elif name in {"ul", "ol"}:
            style = "List Number" if name == "ol" else "List Bullet"
            for li in node.find_all("li", recursive=False):
                paragraph = doc.add_paragraph(li.get_text(" ", strip=True), style=style)
                paragraph.alignment = alignment
            continue
        else:
            paragraph = doc.add_paragraph(text)
        paragraph.alignment = alignment
    doc.core_properties.title = meta["title"]
    doc.core_properties.subject = f"Between Potential and Ideal - {meta['language_label']} edition"
    doc.core_properties.comments = (
        f"Generated from {meta['source_document']}; semantic cross-check: "
        f"{meta['semantic_cross_check']}"
    )
    doc.save(output)


def register_pdf_font() -> str:
    for candidate in FONT_CANDIDATES:
        if candidate.is_file():
            pdfmetrics.registerFont(TTFont("BPIUnicode", str(candidate)))
            return "BPIUnicode"
    return "Helvetica"


def build_pdf(meta: dict[str, str], body: str, output: Path) -> None:
    font = register_pdf_font()
    align = TA_RIGHT if meta["direction"] == "rtl" else TA_LEFT
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "BPIText",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=11,
        leading=17,
        alignment=align,
    )
    heading = ParagraphStyle(
        "BPIHeading",
        parent=normal,
        fontSize=18,
        leading=23,
        textColor="#9b7417",
        spaceBefore=10,
        spaceAfter=8,
    )
    story = [Paragraph(html.escape(meta["title"]), heading), Spacer(1, 5 * mm)]
    rendered = markdown(body, extensions=["extra", "sane_lists"])
    soup = BeautifulSoup(rendered, "html.parser")
    for node in soup.find_all(recursive=False):
        text = node.get_text(" ", strip=True)
        if not text:
            continue
        style = heading if node.name in {"h1", "h2", "h3"} else normal
        story.extend([Paragraph(html.escape(text), style), Spacer(1, 3 * mm)])
    doc = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=meta["title"],
        author="Barak Ben Hur",
        subject=f"Between Potential and Ideal - {meta['language_label']} edition",
    )
    doc.build(story)


def main() -> int:
    parser = ArgumentParser(
        description="Build a localized BPI package in HTML, PDF, DOCX, Markdown and TXT."
    )
    parser.add_argument("source", type=Path)
    parser.add_argument("output_stem", type=Path)
    parser.add_argument(
        "--allow-draft",
        action="store_true",
        help="Allow review output outside site/files from a non-approved source.",
    )
    args = parser.parse_args()

    source = args.source.resolve()
    output_stem = args.output_stem.resolve()
    meta, body = parse_source(source)
    ensure_output_policy(meta, output_stem, args.allow_draft)
    output_stem.parent.mkdir(parents=True, exist_ok=True)

    output_stem.with_suffix(".md").write_text(body, encoding="utf-8")
    output_stem.with_suffix(".txt").write_text(plain_text(body), encoding="utf-8")
    output_stem.with_suffix(".html").write_text(html_document(meta, body), encoding="utf-8")
    build_docx(meta, body, output_stem.with_suffix(".docx"))
    build_pdf(meta, body, output_stem.with_suffix(".pdf"))

    for suffix in (".html", ".pdf", ".docx", ".md", ".txt"):
        path = output_stem.with_suffix(suffix)
        if not path.is_file() or path.stat().st_size == 0:
            raise SystemExit(f"Failed to create {path}")
        print(path.relative_to(ROOT) if path.is_relative_to(ROOT) else path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

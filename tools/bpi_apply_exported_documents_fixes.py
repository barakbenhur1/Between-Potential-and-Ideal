#!/usr/bin/env python3
from __future__ import annotations
from pathlib import Path
import html, json, re, shutil, zipfile
from datetime import datetime, timezone

ROOT = Path.cwd()
REPORT = {"started_at": datetime.now(timezone.utc).isoformat(), "text_files_changed": [], "docx_files_changed": [], "pdf_files_changed": [], "created_files": [], "workflows_disabled": [], "warnings": []}

HE_SPACING_FIXES = {
    "ואידיאליכולים": "ואידיאל יכולים", "אידיאליכולים": "אידיאל יכולים",
    "תהיהאידיאלית": "תהיה אידיאלית", "ביצירהאידיאלי": "ביצירה אידיאלי",
    "אידיאליש": "אידיאל יש", "האידיאליכול": "האידיאל יכול", "אידיאליכול": "אידיאל יכול",
    "שהיהאידיאלי": "שהיה אידיאלי", "שנעשהאידיאלי": "שנעשה אידיאלי", "נעשהאידיאלי": "נעשה אידיאלי",
}
SCIENCE_TERMS_RE = re.compile(r"(P\s*vs\.?\s*NP|NP|Gödel|Godel|גדל|טיורינג|Turing|black hole|חור שחור|event horizon|אופק אירועים|physics|פיזיקה|mathematics|מתמטיקה|infinity|אינסוף|entropy|אנטרופיה|quantum|קוונט|AI|בינה מלאכותית)", re.I)
HE_METHOD_NOTE_MD = """---

**הבהרת מתודולוגיה קצרה:** כאשר המסמך משתמש במונחים מתחומי מתמטיקה, פיזיקה, מדעי המחשב, לוגיקה או בינה מלאכותית, יש לקרוא אותם לפי ההקשר המוצהר: לעיתים כמטאפורה מבנית, לעיתים כמודל קריאה, ולעיתים כטענה פורמלית רק אם הדבר נאמר במפורש. אין להבין שימוש במונח מדעי כהוכחה מדעית או מטפיזית בפני עצמו.
"""
EN_METHOD_NOTE_MD = """---

**Short methodological clarification:** When this document uses terms from mathematics, physics, computer science, logic, or artificial intelligence, they should be read according to their stated context: sometimes as structural metaphor, sometimes as a reading model, and only as a formal claim when explicitly presented as such. The use of a scientific term should not be treated as a scientific or metaphysical proof by itself.
"""
HE_METHOD_NOTE_TXT = HE_METHOD_NOTE_MD.replace("**", "")
EN_METHOD_NOTE_TXT = EN_METHOD_NOTE_MD.replace("**", "")
EDITORIAL_REPORT_EN_MD = '''# Editorial Report - Principles

## A. What was done in the rewritten version

1. Absolute declarations were turned into modal formulations: instead of saying "existence is," "God is," or "the human being is," the wording now relies more on phrases such as "one can think of," "according to the model," "the theory proposes," and "as a way of reading."
2. The core of the theory was preserved: Potential, Ideal, Optimal, distance, suffering, testimony, source, and worthy intelligence.
3. Explicit philosophical humility was added: the text states that it is a thought model rather than a doctrine, and that one may reject the metaphysics while still examining the ethical value.
4. Definitions were sharpened so that the reader does not feel that large concepts are thrown at them without an anchor.
5. Logical bridges were added between the river image, the language about God, the question of suffering, and the question of artificial intelligence.
6. Formulations that sounded like revelation or final proclamation were reduced, while the imagery and poetic force were preserved.
7. The text now stresses that suffering is not good in itself, and that the model must not be used to justify pain, passivity, or external judgment of another person.

## B. Places where there is still a logical leap or a weak claim

1. The transition from the psychology of a single person to a cosmic structure still requires further justification. At present it is presented as an interpretive extension, not as proof.
2. The concept of "the whole" still needs a more precise definition if the text is intended for a strict philosophical setting.
3. The claim that "the whole needs the experience of limitation" is very strong. In the tightened version it was softened into a model, but it still needs to explain why it is not merely a beautiful image.
4. The measurement of "distance" remains an open problem. The text should clarify whether this is a moral, psychological, existential, or metaphysical category.
5. The connection between artificial intelligence and the general metaphysics is interesting, but it still needs one more link: why does a model about existence and suffering lead specifically to a model of "testimony" in intelligence?
6. There is a risk that readers will think the theory romanticizes suffering. It should continue to stress that suffering is not a value, but a condition through which value is sometimes clarified.
7. The language about "God" will provoke resistance. It is recommended to introduce it from the beginning as a technical or metaphorical concept, not as a religious proof.

## C. Suggestions for further strengthening before serious publication

1. Open the essay with a short section titled "What I am not claiming" in order to lower premature resistance.
2. Add a short section titled "A possible secular reading": how the entire model can be understood without God.
3. Add a short section titled "A possible religious reading": how the model can be understood without turning it into religious dogma.
4. Add one concrete example for each concept: Potential, Ideal, Optimal, distance, and testimony.
5. Separate "the philosophical essay" from "the poetic manifesto." The site can contain both, but a serious reader needs to know what is the claim and what is the image.
6. Add a chapter of objections: "Why is this not simply theodicy?" "Why is this not a justification of suffering?" "Why is this not mysticism?" "Why is AI related to this?"
7. Add a very short forum-ready summary: 250-400 words, without the site, without images, and with an explicit invitation to criticism.

## D. Clarification on terminology

I ask in the clearest possible way that the terminology I use in the theory not be treated as correct, precise, final, or binding in itself. These are working names, descriptions, images, and linguistic tools born from a personal attempt to formulate something that is still under examination.

As long as I am alive, I do not agree that another person should use these terms or descriptions as a basis for citation, expansion, method, continuation, or binding interpretation without my explicit signed approval. Anyone who wants to address the idea, criticize it, expand it, or use it as a point of departure is asked not to adopt my terminology automatically, but to formulate things independently and with full responsibility.

After my death, my request remains simpler but no less forceful: think very, very, very carefully before continuing what I tried to do here. If you choose to do so anyway, ask yourself whether the terminology you are using is truly the most precise possible. And even if it seems so, ask again where your own creativity is. If you have a better, more honest, and more precise way to say it, use it. Do not continue my language out of convenience, admiration, or laziness.

And if you are foolish enough to become attached to what I tried to say here, and, God forbid, to find sense in it, or even a small edge of truth, at least be free enough to say it in your own way.

## E. Personal closing note

I feel that even if new ideas or further refinements to the theory come to me, I probably will not add them anymore.

*I realized that I need to walk my dog.*
'''

def is_hebrew_path(p: Path) -> bool:
    s = p.as_posix().lower()
    return any(x in s for x in ["-he.", "-hebrew", "/he/", "rtl", "-he-"])

def fix_text_content(text: str) -> str:
    for bad, good in HE_SPACING_FIXES.items():
        text = text.replace(bad, good)
    return text

def needs_method(path: Path, text: str) -> bool:
    s = path.as_posix().lower()
    if any(skip in s for skip in ["appendices/stories", "mistake-repeats", "readme"]): return False
    return any(x in s for x in ["between-potential", "editorial-tightened", "potential-extensions", "ai-believes"]) and (SCIENCE_TERMS_RE.search(text) or "between-potential" in s or "editorial-tightened" in s)

def add_method_note(path: Path, text: str, markdown: bool) -> str:
    if "Short methodological clarification" in text or "הבהרת מתודולוגיה קצרה" in text: return text
    note = (HE_METHOD_NOTE_MD if markdown else HE_METHOD_NOTE_TXT) if is_hebrew_path(path) else (EN_METHOD_NOTE_MD if markdown else EN_METHOD_NOTE_TXT)
    return text.rstrip()+"\n\n"+note+"\n"

def patch_md_txt():
    for path in list((ROOT/"site").rglob("*.md"))+list((ROOT/"site").rglob("*.txt")):
        try: text = path.read_text(encoding="utf-8")
        except Exception: continue
        old = text; text = fix_text_content(text)
        if needs_method(path, text): text = add_method_note(path, text, path.suffix.lower()==".md")
        if text != old:
            path.write_text(text, encoding="utf-8"); REPORT["text_files_changed"].append(path.as_posix())

def create_editorial_en():
    base = ROOT/"site/files/editorial-tightened"; base.mkdir(parents=True, exist_ok=True)
    for name, content in {"editorial-report-en.md": EDITORIAL_REPORT_EN_MD, "editorial-report-en.txt": re.sub(r"[#*_`>-]", "", EDITORIAL_REPORT_EN_MD)}.items():
        p = base/name
        if not p.exists(): p.write_text(content, encoding="utf-8"); REPORT["created_files"].append(p.as_posix())
    htmlp = base/"editorial-report-en.html"
    if not htmlp.exists():
        body=[]
        for line in EDITORIAL_REPORT_EN_MD.splitlines():
            if line.startswith("# "): body.append(f"<h1>{html.escape(line[2:])}</h1>")
            elif line.startswith("## "): body.append(f"<h2>{html.escape(line[3:])}</h2>")
            elif line.strip(): body.append(f"<p>{html.escape(line)}</p>")
        htmlp.write_text(f'<!DOCTYPE html>\n<html dir="ltr" lang="en"><head><meta charset="utf-8"/><title>Editorial Report - Principles</title><link href="../../styles.css" rel="stylesheet"/></head><body><main>{"".join(body)}</main></body></html>\n', encoding="utf-8")
        REPORT["created_files"].append(htmlp.as_posix())
    docxp = base/"editorial-report-en.docx"
    if not docxp.exists():
        try:
            from docx import Document
            d=Document()
            for line in EDITORIAL_REPORT_EN_MD.splitlines():
                if line.startswith("# "): d.add_heading(line[2:], level=1)
                elif line.startswith("## "): d.add_heading(line[3:], level=2)
                elif line.strip(): d.add_paragraph(line)
            d.save(docxp); REPORT["created_files"].append(docxp.as_posix())
        except Exception as e: REPORT["warnings"].append(f"Could not create {docxp}: {e}")
    pdfp = base/"editorial-report-en.pdf"
    if not pdfp.exists():
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import cm
            styles=getSampleStyleSheet(); story=[]
            for line in EDITORIAL_REPORT_EN_MD.splitlines():
                if line.startswith("# "): story += [Paragraph(html.escape(line[2:]), styles["Title"]), Spacer(1,.35*cm)]
                elif line.startswith("## "): story += [Paragraph(html.escape(line[3:]), styles["Heading2"]), Spacer(1,.2*cm)]
                elif line.strip(): story += [Paragraph(html.escape(line), styles["BodyText"]), Spacer(1,.12*cm)]
            SimpleDocTemplate(str(pdfp), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm).build(story)
            REPORT["created_files"].append(pdfp.as_posix())
        except Exception as e: REPORT["warnings"].append(f"Could not create {pdfp}: {e}")

def alt_for_path(p: Path, idx: int) -> str:
    name = p.stem.replace("-", " ").replace("_", " ")
    return f"איור במסמך: {name}, תמונה {idx}" if is_hebrew_path(p) else f"Figure in document: {name}, image {idx}"

def add_docx_note(xml: str, p: Path) -> str:
    if "Short methodological clarification" in xml or "הבהרת מתודולוגיה קצרה" in xml: return xml
    s=p.as_posix().lower()
    if not any(x in s for x in ["between-potential", "editorial-tightened", "potential-extensions"]): return xml
    note = HE_METHOD_NOTE_TXT if is_hebrew_path(p) else EN_METHOD_NOTE_TXT
    esc = html.escape(note)
    para = '<w:p><w:r><w:t xml:space="preserve">---</w:t></w:r></w:p>' + f'<w:p><w:r><w:t xml:space="preserve">{esc}</w:t></w:r></w:p>'
    idx=xml.rfind("<w:sectPr")
    return xml[:idx]+para+xml[idx:] if idx!=-1 else xml.replace("</w:body>", para+"</w:body>")

def patch_docx(p: Path):
    tmp=p.with_suffix(p.suffix+".tmp"); changed=False
    try:
        with zipfile.ZipFile(p,"r") as zin, zipfile.ZipFile(tmp,"w",zipfile.ZIP_DEFLATED) as zout:
            img_i=0
            for name in zin.namelist():
                data=zin.read(name)
                if name=="word/document.xml":
                    xml=data.decode("utf-8"); old=xml; xml=fix_text_content(xml)
                    def repl(m):
                        nonlocal img_i
                        tag=m.group(0); img_i+=1
                        if re.search(r'\sdescr="[^"]+"', tag): return tag
                        alt=html.escape(alt_for_path(p,img_i), quote=True)
                        return tag[:-2]+f' descr="{alt}" title="{alt}"/>'
                    xml=re.sub(r'<wp:docPr\b[^>]*/>', repl, xml)
                    def repl2(m):
                        tag=m.group(0)
                        if re.search(r'\sdescr="[^"]+"', tag): return tag
                        alt=html.escape(alt_for_path(p,1), quote=True)
                        return tag[:-2]+f' descr="{alt}"/>'
                    xml=re.sub(r'<pic:cNvPr\b[^>]*/>', repl2, xml)
                    xml=add_docx_note(xml,p)
                    if xml!=old: changed=True
                    data=xml.encode("utf-8")
                elif name=="docProps/core.xml":
                    xml=data.decode("utf-8"); old=xml
                    xml=xml.replace("<dc:creator>python-docx</dc:creator>","<dc:creator>Barak Ben Hur</dc:creator>")
                    if xml!=old: changed=True
                    data=xml.encode("utf-8")
                zout.writestr(name,data)
        if changed:
            shutil.move(tmp,p); REPORT["docx_files_changed"].append(p.as_posix())
        else: tmp.unlink(missing_ok=True)
    except Exception as e:
        tmp.unlink(missing_ok=True); REPORT["warnings"].append(f"Could not patch DOCX {p}: {e}")

def patch_docx_all():
    for p in (ROOT/"site").rglob("*.docx"): patch_docx(p)

def patch_pdf_meta():
    try: import fitz
    except Exception as e: REPORT["warnings"].append(f"PyMuPDF unavailable: {e}"); return
    for p in (ROOT/"site").rglob("*.pdf"):
        try:
            doc=fitz.open(p); meta=doc.metadata or {}; old=dict(meta)
            title=p.stem.replace("-"," ").replace("_"," ").strip().title()
            meta["title"] = meta.get("title") or title
            meta["author"] = meta.get("author") or "Barak Ben Hur"
            meta["subject"] = meta.get("subject") or "Between Potential and Ideal project document"
            meta["keywords"] = meta.get("keywords") or "Between Potential and Ideal, Potential, Ideal, philosophy, theory"
            if meta!=old:
                doc.set_metadata(meta); tmp=p.with_suffix(".pdf.tmp"); doc.save(tmp, incremental=False, deflate=True, garbage=3); doc.close(); shutil.move(tmp,p); REPORT["pdf_files_changed"].append(p.as_posix())
            else: doc.close()
        except Exception as e: REPORT["warnings"].append(f"Could not update PDF metadata {p}: {e}")

def disable_workflows():
    wf=ROOT/".github/workflows"
    if not wf.exists(): return
    for p in wf.glob("*.yml"):
        if not any(k in p.name for k in ["one-time", "fix-", "bpi-v86", "bpi-v87", "bpi-v88"]): continue
        text=p.read_text(encoding="utf-8"); old=text
        text=re.sub(r"on:\n(?:  workflow_dispatch:\n)?  push:\n    branches: \[main\]\n    paths:\n(?:      - .*\n)+", "on:\n  workflow_dispatch:\n", text)
        text=re.sub(r"on:\n  push:\n    branches: \[main\]\n    paths:\n(?:      - .*\n)+  workflow_dispatch:\n", "on:\n  workflow_dispatch:\n", text)
        if text!=old:
            p.write_text("# Push trigger disabled by BPI ideal fixer to prevent one-time workflow regressions.\n"+text, encoding="utf-8")
            REPORT["workflows_disabled"].append(p.as_posix())

def write_reports():
    REPORT["finished_at"] = datetime.now(timezone.utc).isoformat()
    (ROOT/"BPI_EXPORTED_DOCUMENTS_FIX_REPORT.json").write_text(json.dumps(REPORT, ensure_ascii=False, indent=2), encoding="utf-8")
    md=["# BPI Exported Documents Fix Report", "", f"- Text files changed: {len(REPORT['text_files_changed'])}", f"- DOCX files changed: {len(REPORT['docx_files_changed'])}", f"- PDF metadata changed: {len(REPORT['pdf_files_changed'])}", f"- Files created: {len(REPORT['created_files'])}", f"- Workflows with push disabled: {len(REPORT['workflows_disabled'])}", f"- Warnings: {len(REPORT['warnings'])}", "", "## Created files"]
    md += [f"- `{x}`" for x in REPORT["created_files"]]
    md += ["", "## Important limitation", "`site/files/appendices/mistake-repeats/infinity-pool-original-he.pdf` is an image-only Hebrew PDF in the provided export set. No faithful English counterpart was generated from it because the source text was not available. Creating a high-quality translation requires the original text source or explicit approval to OCR and manually review it.", "", "## Warnings"]
    md += [f"- {w}" for w in REPORT["warnings"]]
    (ROOT/"docs/BPI_EXPORTED_DOCUMENTS_FIX_REPORT.md").write_text("\n".join(md)+"\n", encoding="utf-8")

def main():
    patch_md_txt(); create_editorial_en(); patch_docx_all(); patch_pdf_meta(); disable_workflows(); write_reports(); print(json.dumps(REPORT, ensure_ascii=False, indent=2))
if __name__ == "__main__": main()

#!/usr/bin/env python3
"""Build public beta packages for the complete Klingon and Neo-Quenya editions."""

from __future__ import annotations

import hashlib
import html
import json
import re
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
CONTRACT = ROOT / "localization/documents/between-potential-and-ideal.json"
SITE = ROOT / "site"
BETA_MANIFEST = ROOT / "localization/beta-release-manifest.json"
LANGUAGES = {
    "tlh": {
        "label": "tlhIngan Hol",
        "html_lang": "tlh",
        "notice": "Public beta translation. Klingon linguistic review is still active; this edition is not presented as canonical Klingon.",
        "title": "Between Potential and Ideal — tlhIngan Hol Public Beta",
        "intro": "ghItlh naQ mughlu'pu'bogh laDlaHlu'. mughghachvam beta 'oH; Hol po'wI' nuDghach taH.",
    },
    "qya": {
        "label": "Neo-Quenya",
        "html_lang": "qya",
        "notice": "Public beta translation. Neo-Quenya linguistic review is still active; this edition is explicitly a modern reconstruction, not Tolkien-authored text.",
        "title": "Between Potential and Ideal — Neo-Quenya Public Beta",
        "intro": "I quentalë quanta ná sí laitanwa ve public beta. I lambë ná Neo-Quenya, ar i metta parmaquetalië lemya carienna.",
    },
}
CONTROL_HEADINGS = {
    "## Segment review gate",
    "## Placeholder review gate",
}
BASE_CONTROL_MARKERS = (
    "Translation control note",
    "mIw qawmoHghach",
    "Enyalë Léo",
)


def strip_front_matter(text: str, path: Path) -> str:
    if not text.startswith("---\n"):
        raise RuntimeError(f"Missing front matter: {path.relative_to(ROOT)}")
    parts = text.split("---\n", 2)
    if len(parts) != 3:
        raise RuntimeError(f"Malformed front matter: {path.relative_to(ROOT)}")
    return parts[2].strip()


def strip_review_gate(text: str) -> str:
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if line.strip() in CONTROL_HEADINGS:
            return "\n".join(lines[:index]).rstrip()
    return text.strip()


def assemble(language: str, contract: dict) -> str:
    base_path = ROOT / contract["canonical_targets"][language]
    base = strip_front_matter(base_path.read_text(encoding="utf-8"), base_path)
    for marker in BASE_CONTROL_MARKERS:
        needle = f"\n## {marker}"
        position = base.find(needle)
        if position >= 0:
            base = base[:position].rstrip()
            break

    parts = [base]
    for relative in contract["source_segments"][language]:
        path = ROOT / relative
        body = strip_front_matter(path.read_text(encoding="utf-8"), path)
        parts.append(strip_review_gate(body))

    notice = LANGUAGES[language]["notice"]
    return f"> **Public Beta — linguistic review ongoing.** {notice}\n\n" + "\n\n".join(
        item.strip() for item in parts if item.strip()
    ) + "\n"


def markdown_to_plain(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"[Image: \1]", text)
    text = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*+]\s+", "• ", text, flags=re.MULTILINE)
    text = text.replace("**", "").replace("__", "").replace("*", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return html.unescape(text).strip() + "\n"


def render_html(language: str, markdown_text: str, stem: str) -> str:
    info = LANGUAGES[language]
    body = markdown.markdown(
        markdown_text,
        extensions=["extra", "tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )
    downloads = " ".join(
        f'<a href="{stem}.{ext}">{ext.upper()}</a>' for ext in ("pdf", "docx", "md", "txt")
    )
    return f"""<!doctype html>
<html lang="{info['html_lang']}" dir="ltr">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html.escape(info['title'])}</title>
<meta name="description" content="Complete public beta translation of Between Potential and Ideal in {html.escape(info['label'])}.">
<meta name="robots" content="index,follow">
<style>
:root{{color-scheme:dark}}body{{margin:0;background:#080b12;color:#edf2ff;font:18px/1.72 system-ui,-apple-system,sans-serif}}main{{max-width:920px;margin:auto;padding:32px 22px 80px}}a{{color:#7fd7ff}}h1,h2,h3{{line-height:1.2;color:#fff;margin-top:2em}}blockquote{{margin:0 0 2rem;padding:1rem 1.25rem;border:2px solid #f2c45e;background:#211b0e;color:#fff7db}}nav{{position:sticky;top:0;background:#080b12ee;padding:12px 0;border-bottom:1px solid #2d3545;z-index:2}}nav a{{margin-right:14px;font-weight:700}}img{{max-width:100%;height:auto}}code{{white-space:pre-wrap}}@media(max-width:600px){{body{{font-size:16px}}main{{padding:20px 16px 60px}}}}
</style>
</head>
<body><main>
<nav><a href="../../{language}.html">Beta home</a> {downloads} <a href="../../en.html">English</a> <a href="../../index.html">עברית</a></nav>
{body}
</main></body></html>
"""


def build_docx(language: str, text: str, output: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "DejaVu Sans"
    normal.font.size = Pt(10.5)
    doc.core_properties.title = LANGUAGES[language]["title"]
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            doc.add_heading(re.sub(r"[*_`]", "", heading.group(2)), level=min(len(heading.group(1)), 4))
        elif line.startswith(">"):
            doc.add_paragraph(re.sub(r"^>\s*", "", line).replace("**", ""), style="Quote")
        elif re.match(r"^[-*+]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*+]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        elif line.startswith("!["):
            continue
        else:
            doc.add_paragraph(re.sub(r"[*_`]", "", line))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def register_pdf_font() -> str:
    candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
    )
    for candidate in candidates:
        if candidate.is_file():
            pdfmetrics.registerFont(TTFont("BetaUnicode", str(candidate)))
            return "BetaUnicode"
    return "Helvetica"


def build_pdf(language: str, text: str, output: Path) -> None:
    font = register_pdf_font()
    styles = getSampleStyleSheet()
    base = ParagraphStyle("BetaBody", parent=styles["BodyText"], fontName=font, fontSize=9.2, leading=13.2, spaceAfter=6)
    heading = ParagraphStyle("BetaHeading", parent=base, fontSize=15, leading=18, spaceBefore=12, spaceAfter=8)
    title = ParagraphStyle("BetaTitle", parent=heading, fontSize=20, leading=24, alignment=TA_CENTER)
    notice = ParagraphStyle("BetaNotice", parent=base, fontSize=9, leading=12, borderWidth=1, borderPadding=7, spaceAfter=12)
    story = []
    first_title = True
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 2 * mm))
            continue
        if line.startswith("!["):
            continue
        clean = markdown_to_plain(line).strip()
        if not clean:
            continue
        escaped = html.escape(clean).replace("\n", "<br/>")
        if line.startswith(">"):
            story.append(Paragraph(escaped, notice))
        elif line.startswith("#"):
            story.append(Paragraph(escaped, title if first_title else heading))
            first_title = False
        else:
            story.append(Paragraph(escaped, base))
    output.parent.mkdir(parents=True, exist_ok=True)
    pdf = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=17 * mm, bottomMargin=17 * mm, title=LANGUAGES[language]["title"], author="Barak Ben Hur")
    pdf.build(story or [Paragraph("Empty document", base)])


def gateway(language: str, stem: str) -> str:
    info = LANGUAGES[language]
    buttons = "".join(
        f'<a class="button" href="files/{language}/{stem}.{ext}">{ext.upper()}</a>'
        for ext in ("html", "pdf", "docx", "md", "txt")
    )
    return f"""<!doctype html><html lang="{info['html_lang']}" dir="ltr"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>{html.escape(info['title'])}</title><meta name="description" content="Complete public beta edition in {html.escape(info['label'])}."><style>:root{{color-scheme:dark}}body{{margin:0;background:#070a10;color:#f2f5ff;font:18px/1.6 system-ui,sans-serif;display:grid;min-height:100vh;place-items:center}}main{{max-width:780px;padding:34px;text-align:center}}.beta{{display:inline-block;padding:.35rem .7rem;border:1px solid #f2c45e;color:#f2c45e;border-radius:999px;font-weight:800}}h1{{font-size:clamp(2rem,6vw,4rem);line-height:1.05}}.button{{display:inline-block;margin:.4rem;padding:.75rem 1rem;border-radius:.6rem;background:#132039;color:#8edcff;text-decoration:none;font-weight:800}}.warning{{margin:1.5rem 0;padding:1rem;border:2px solid #f2c45e;background:#211b0e}}a{{color:#8edcff}}</style></head><body><main><span class="beta">PUBLIC BETA</span><h1>{html.escape(info['label'])}</h1><p>{html.escape(info['intro'])}</p><p class="warning">{html.escape(info['notice'])}</p><div>{buttons}</div><p><a href="en.html">English</a> · <a href="index.html">עברית</a></p></main></body></html>"""


def digest(path: Path) -> dict:
    data = path.read_bytes()
    return {"path": path.relative_to(ROOT).as_posix(), "bytes": len(data), "sha256": hashlib.sha256(data).hexdigest()}


def main() -> int:
    contract = json.loads(CONTRACT.read_text(encoding="utf-8"))
    files = []
    for language in LANGUAGES:
        text = assemble(language, contract)
        folder = SITE / "files" / language
        folder.mkdir(parents=True, exist_ok=True)
        stem = f"between-potential-and-ideal-{language}"
        md_path = folder / f"{stem}.md"
        txt_path = folder / f"{stem}.txt"
        html_path = folder / f"{stem}.html"
        docx_path = folder / f"{stem}.docx"
        pdf_path = folder / f"{stem}.pdf"
        md_path.write_text(text, encoding="utf-8")
        txt_path.write_text(markdown_to_plain(text), encoding="utf-8")
        html_path.write_text(render_html(language, text, stem), encoding="utf-8")
        build_docx(language, text, docx_path)
        build_pdf(language, text, pdf_path)
        (SITE / f"{language}.html").write_text(gateway(language, stem), encoding="utf-8")
        generated = [md_path, txt_path, html_path, docx_path, pdf_path, SITE / f"{language}.html"]
        for path in generated:
            if not path.is_file() or path.stat().st_size == 0:
                raise RuntimeError(f"Beta output is missing or empty: {path.relative_to(ROOT)}")
            files.append(digest(path))

    manifest = {
        "schema_version": 1,
        "release_channel": "public-beta",
        "linguistic_review_complete": False,
        "languages": list(LANGUAGES),
        "segments_per_language": len(contract["source_segments"]["tlh"]),
        "files": files,
    }
    BETA_MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Built {len(files)} public beta files from {manifest['segments_per_language']} segments per language.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
